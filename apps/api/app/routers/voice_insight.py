import logging
import uuid
import json
import re
from datetime import datetime, timedelta, timezone

from fastapi import APIRouter, Depends, HTTPException, Query, Request, UploadFile, File
from fastapi.responses import JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func

from app.dependencies import get_db, get_current_user
from app.models.call_analysis import CallAnalysis, AnalysisStatus
from app.schemas.voice_insight import (
    CallAnalysisCreate,
    CallAnalysisResponse,
    CallAnalysisDetailResponse,
    CallAnalysisListResponse,
)
from app.services import clone_service
from app.services.voice_insight_service import (
    start_gladia_transcription,
    upload_audio_to_gladia,
    get_gladia_transcription,
)
from app.services.storage import upload_file as upload_to_r2, get_download_presigned_url
from app.exceptions import ExternalServiceError
from app.config import settings

logger = logging.getLogger("tarang.voice_insight")

# Leave sufficient context for the instruction file and a complete JSON
# response. Without this, a long Gladia transcript can fail before inference.
_MAX_TRANSCRIPT_PROMPT_CHARS = 7_500
_SARVAM_MAX_TOKENS = 2_048

router = APIRouter(
    prefix="/api/v1/voice-insight",
    tags=["voice-insight"],
)



# ─────────────────────────────────────────────────────────────────────────────
# Load VoiceInsight instruction file at module startup (cached, not per-request)
# ─────────────────────────────────────────────────────────────────────────────
import pathlib as _pathlib

_INSTRUCTION_PATH = _pathlib.Path(__file__).resolve().parent.parent / "instructions" / "voice_insight_instruction.md"
_INSTRUCTION_TEXT = ""
if _INSTRUCTION_PATH.exists():
    _INSTRUCTION_TEXT = _INSTRUCTION_PATH.read_text(encoding="utf-8")
    logger.info("Loaded VoiceInsight instruction file (%d chars)", len(_INSTRUCTION_TEXT))
else:
    logger.warning("VoiceInsight instruction file not found at %s", _INSTRUCTION_PATH)


def _format_timestamp(seconds: float) -> str:
    """Convert seconds to MM:SS format."""
    minutes = int(seconds) // 60
    secs = int(seconds) % 60
    return f"{minutes:02d}:{secs:02d}"


def _build_structured_transcript(transcript_data: dict) -> tuple[str, dict]:
    """
    Build a structured transcript with timestamps from Gladia V2 response.
    Returns (formatted_text, gladia_extras) where gladia_extras contains
    sentiment and NER data for injection into the prompt.
    """
    res_obj = transcript_data.get("result", transcript_data.get("prediction", transcript_data))
    utterances = (
        res_obj.get("utterances")
        or (res_obj.get("transcription", {}).get("utterances") if isinstance(res_obj.get("transcription"), dict) else None)
        or []
    )

    # Extract Gladia-provided sentiment and NER data
    gladia_extras = {}
    if isinstance(res_obj.get("transcription"), dict):
        transcription_obj = res_obj["transcription"]
        if "sentiment_analysis" in transcription_obj:
            gladia_extras["sentiment_analysis"] = transcription_obj["sentiment_analysis"]
        if "named_entity_recognition" in transcription_obj:
            gladia_extras["named_entity_recognition"] = transcription_obj["named_entity_recognition"]
    # Also check at top-level result
    if "sentiment_analysis" in res_obj:
        gladia_extras["sentiment_analysis"] = res_obj["sentiment_analysis"]
    if "named_entity_recognition" in res_obj:
        gladia_extras["named_entity_recognition"] = res_obj["named_entity_recognition"]

    if not utterances:
        # Fallback: plain text transcript
        if isinstance(res_obj.get("transcription"), str):
            return res_obj["transcription"], gladia_extras
        elif isinstance(res_obj.get("full_transcript"), str):
            return res_obj["full_transcript"], gladia_extras
        elif isinstance(res_obj.get("transcription"), dict):
            ft = res_obj["transcription"].get("full_transcript", "")
            if ft:
                return ft, gladia_extras
        return str(res_obj), gladia_extras

    # Build structured transcript with timestamps
    lines = []
    for utt in utterances:
        start = utt.get("start", 0)
        speaker = utt.get("speaker", "Unknown")
        text = utt.get("text", "")
        ts = _format_timestamp(start)
        lines.append(f"[{ts}] Speaker {speaker}: {text}")

    return "\n".join(lines), gladia_extras


def _fit_transcript_for_prompt(transcript: str) -> str:
    """Keep long calls within the deployed Modal model's context window."""
    if len(transcript) <= _MAX_TRANSCRIPT_PROMPT_CHARS:
        return transcript

    marker = "\n\n[Middle of transcript omitted to fit the analysis window]\n\n"
    head_length = int((_MAX_TRANSCRIPT_PROMPT_CHARS - len(marker)) * 0.65)
    tail_length = _MAX_TRANSCRIPT_PROMPT_CHARS - len(marker) - head_length
    return f"{transcript[:head_length]}{marker}{transcript[-tail_length:]}"


def _json_candidates(raw: str):
    """Yield complete JSON-object candidates without being confused by braces in strings."""
    for start in (match.start() for match in re.finditer(r"\{", raw)):
        depth = 0
        in_string = False
        escaped = False
        for index in range(start, len(raw)):
            char = raw[index]
            if in_string:
                if escaped:
                    escaped = False
                elif char == "\\":
                    escaped = True
                elif char == '"':
                    in_string = False
                continue
            if char == '"':
                in_string = True
            elif char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    yield raw[start:index + 1]
                    break


def _parse_sarvam_json(response_text: str) -> dict:
    """Recover the report JSON from common model wrappers and minor JSON defects."""
    raw = re.sub(r"<think>.*?</think>", "", response_text, flags=re.DOTALL).strip()
    raw = re.sub(r"```(?:json)?\s*", "", raw, flags=re.IGNORECASE).replace("```", "")

    candidates = list(_json_candidates(raw))
    for candidate in candidates:
        for attempt in (candidate, re.sub(r",\s*([}\]])", r"\1", candidate)):
            try:
                decoded = json.loads(attempt)
                if isinstance(decoded, dict):
                    return decoded
            except json.JSONDecodeError:
                continue
    raise json.JSONDecodeError("No valid JSON object in Sarvam response", raw, 0)


# ─────────────────────────────────────────────────────────────────────────────
# Synchronous Intelligence Extraction (Sarvam-30B FP8 via Modal)
# ─────────────────────────────────────────────────────────────────────────────
# WHY SYNC: BackgroundTasks and asyncio.create_task() die in Vercel/Cloud Run
# serverless. Instead, Sarvam extraction runs inline within the HTTP request.

async def _run_sarvam_extraction_sync(transcript_data: dict) -> dict:
    """Call Sarvam-30B on Modal synchronously and return the intelligence JSON.
    
    Returns dict with intelligence data or {"error": "..."} on failure.
    No DB writes — caller handles persistence.
    """
    import httpx

    # Build structured transcript with timestamps
    conversation_text, gladia_extras = _build_structured_transcript(transcript_data)
    conversation_text = _fit_transcript_for_prompt(conversation_text)

    # Build the analysis prompt with instruction file + transcript + Gladia extras
    extra_context = ""
    if gladia_extras.get("sentiment_analysis"):
        extra_context += f"\n\n## Gladia Sentiment Analysis Data (use as reference):\n{json.dumps(gladia_extras['sentiment_analysis'], indent=2, ensure_ascii=False)[:2000]}"
    if gladia_extras.get("named_entity_recognition"):
        extra_context += f"\n\n## Gladia NER Data (use as reference):\n{json.dumps(gladia_extras['named_entity_recognition'], indent=2, ensure_ascii=False)[:2000]}"

    user_prompt = (
        f"Analyze the following call recording transcript.\n\n"
        f"## Transcript:\n{conversation_text}\n"
        f"{extra_context}\n\n"
        f"Now analyze this transcript and output the structured JSON as specified in your instructions. "
        f"Keep arrays concise so the entire response remains valid JSON."
    )

    messages = [
        {"role": "system", "content": _INSTRUCTION_TEXT or "You are a law enforcement intelligence analyst. Output strict JSON only."},
        {"role": "user", "content": user_prompt},
    ]

    try:
        if not settings.MODAL_SARVAM_INSIGHT_ENDPOINT:
            raise ValueError("MODAL_SARVAM_INSIGHT_ENDPOINT not configured")

        async with httpx.AsyncClient(timeout=300.0, follow_redirects=True) as client:
            headers = {"Content-Type": "application/json"}
            # Avoid sending an empty credential in local development.
            if settings.MODAL_SHARED_SECRET:
                headers["x-tarang-modal-secret"] = settings.MODAL_SHARED_SECRET
            body = {
                "messages": messages,
                "temperature": 0.1,
                "max_tokens": _SARVAM_MAX_TOKENS,
                "response_format": {"type": "json_object"},
            }
            resp = await client.post(settings.MODAL_SARVAM_INSIGHT_ENDPOINT, headers=headers, json=body)
            resp.raise_for_status()
            response_data = resp.json()
            response_text = response_data.get("content", "")
            if not isinstance(response_text, str) or not response_text.strip():
                raise ValueError("Sarvam returned an empty response")
            response_text = response_text.strip()

        intelligence_json = _parse_sarvam_json(response_text)
        logger.info("[OK] Sarvam-30B intelligence extracted successfully")
        return intelligence_json

    except json.JSONDecodeError:
        logger.error("Sarvam returned non-JSON: %s", response_text[:300] if 'response_text' in locals() else "N/A")
        return {"error": "Non-JSON response from LLM"}
    except httpx.TimeoutException:
        logger.error("Sarvam extraction timed out")
        return {"error": "Sarvam analysis timed out. Please retry the analysis."}
    except httpx.HTTPStatusError as e:
        logger.error("Sarvam endpoint returned HTTP %s: %s", e.response.status_code, e.response.text[:500])
        return {"error": f"Sarvam service returned HTTP {e.response.status_code}. Please retry the analysis."}
    except Exception as e:
        logger.error("Sarvam extraction failed: %s", e)
        return {"error": str(e)}


# ─────────────────────────────────────────────────────────────────────────────
# Endpoints
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/upload-audio")
async def upload_audio_to_r2(
    file: UploadFile = File(...),
    clerk_user_id: str = Depends(get_current_user),
):
    """Upload an audio recording file directly to Cloudflare R2 and Gladia."""
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty audio file provided")

    # 1. Store in Cloudflare R2 for permanent archive
    r2_key = f"voice_insight/{uuid.uuid4()}_{file.filename}"
    upload_to_r2(file_bytes, r2_key, file.content_type or "application/octet-stream")

    # 2. Upload directly to Gladia API to ensure 100% accessible transcription URL
    try:
        gladia_audio_url = await upload_audio_to_gladia(file_bytes, file.filename)
    except Exception as e:
        logger.warning("Gladia direct upload failed, falling back to R2 presigned URL: %s", e)
        gladia_audio_url = get_download_presigned_url(r2_key, expiration=86400)

    return {"r2_key": r2_key, "audio_url": gladia_audio_url, "filename": file.filename}


@router.post("/analyze", response_model=CallAnalysisResponse)
async def analyze_call(
    body: CallAnalysisCreate,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Start analyzing a call recording via Gladia transcription.
    
    Starts the Gladia job and saves the job ID. The frontend polls
    GET /calls/{id} which checks Gladia status inline (self-healing).
    Once transcript is ready, user triggers Sarvam extraction via
    POST /calls/{id}/extract-intelligence.
    
    No background tasks — everything is serverless-safe.
    """
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    try:
        gladia_response = await start_gladia_transcription(
            str(body.audio_url),
            source_language=body.source_language,
            translation=body.translation,
            translation_target_language=body.translation_target_language,
        )
        gladia_job_id = gladia_response.get("id")
    except ExternalServiceError as e:
        logger.error("Gladia service error: %s", e.message)
        raise HTTPException(status_code=502, detail=e.message)
    except Exception as e:
        logger.error("Failed to start Gladia job: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Transcription error: {str(e)}")

    call = CallAnalysis(
        user_id=user_id,
        filename=body.filename,
        audio_url=str(body.audio_url),
        audio_r2_key=body.audio_r2_key,
        status=AnalysisStatus.TRANSCRIBING,
        gladia_job_id=gladia_job_id,
    )
    db.add(call)
    await db.commit()
    await db.refresh(call)

    return call





@router.get("/calls", response_model=CallAnalysisListResponse)
async def list_calls(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    q: str = Query(None, description="Search query"),
    status: str = Query(None, description="Filter by status"),
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """List and search call analysis records for the current user."""
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    offset = (page - 1) * per_page
    query = select(CallAnalysis).where(CallAnalysis.user_id == user_id)

    if q:
        query = query.where(CallAnalysis.filename.ilike(f"%{q}%"))

    if status:
        query = query.where(CallAnalysis.status == status)

    # Count
    count_query = select(func.count()).select_from(query.subquery())
    total = (await db.execute(count_query)).scalar()

    # Paginate
    query = query.order_by(CallAnalysis.created_at.desc()).offset(offset).limit(per_page)
    calls = (await db.execute(query)).scalars().all()

    return {"items": calls, "total": total}


@router.get("/calls/{call_id}", response_model=CallAnalysisDetailResponse)
async def get_call_detail(
    call_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get details for a specific call analysis.

    SELF-HEALING: Checks Gladia status inline on each poll.
    When Gladia is done, saves transcript and sets status to
    TRANSCRIPT_READY so the user can trigger Sarvam extraction.
    No background tasks — fully serverless-safe.
    """
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    result = await db.execute(
        select(CallAnalysis).where(
            CallAnalysis.id == call_id,
            CallAnalysis.user_id == user_id,
        )
    )
    call = result.scalar_one_or_none()
    if not call:
        raise HTTPException(status_code=404, detail="Call analysis not found")

    # -- Self-healing: poll Gladia inline when still transcribing --

    if call.status == AnalysisStatus.TRANSCRIBING and call.gladia_job_id:
        try:
            gladia_res = await get_gladia_transcription(call.gladia_job_id)
            gladia_status = gladia_res.get("status")

            if gladia_status == "done":
                logger.info("[HEAL] Gladia done for call %s — saving transcript", call_id)
                call.transcript = gladia_res
                call.status = AnalysisStatus.TRANSCRIPT_READY

                # Extract audio duration
                res_obj = gladia_res.get("result", gladia_res)
                meta = res_obj.get("metadata", {}) if isinstance(res_obj, dict) else {}
                if meta.get("audio_duration"):
                    call.duration_seconds = float(meta["audio_duration"])

                await db.commit()
                await db.refresh(call)

            elif gladia_status == "error":
                logger.error("[HEAL] Gladia failed for call %s", call_id)
                call.status = AnalysisStatus.FAILED
                await db.commit()
                await db.refresh(call)
        except Exception as e:
            logger.warning("[HEAL] Gladia check failed for call %s: %s", call_id, e)

    # Migrate legacy EXTRACTING records that never got intelligence
    elif call.status == AnalysisStatus.EXTRACTING and not call.intelligence and call.transcript:
        logger.info("[HEAL] Migrating stale EXTRACTING call %s to TRANSCRIPT_READY", call_id)
        call.status = AnalysisStatus.TRANSCRIPT_READY
        await db.commit()
        await db.refresh(call)

    response = CallAnalysisDetailResponse.model_validate(call)
    if call.audio_r2_key:
        response.playback_url = get_download_presigned_url(call.audio_r2_key, expiration=3600)
    return response


@router.post("/calls/{call_id}/extract-intelligence", response_model=CallAnalysisDetailResponse)
async def extract_intelligence(
    call_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Synchronously run Sarvam-30B intelligence extraction on a transcribed call.
    
    This is the user-triggered step that replaces the broken background task.
    Runs entirely within the HTTP request (up to 300s for Modal inference).
    No background tasks, no asyncio.create_task — fully serverless-safe.
    """
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    result = await db.execute(
        select(CallAnalysis).where(
            CallAnalysis.id == call_id,
            CallAnalysis.user_id == user_id,
        )
    )
    call = result.scalar_one_or_none()
    if not call:
        raise HTTPException(status_code=404, detail="Call analysis not found")

    if not call.transcript:
        raise HTTPException(status_code=400, detail="No transcript available — transcription must complete first")

    if call.status == AnalysisStatus.EXTRACTING:
        raise HTTPException(status_code=409, detail="Intelligence extraction is already in progress")

    # Mark as extracting so the UI shows the right state
    call.status = AnalysisStatus.EXTRACTING
    await db.commit()

    # Run Sarvam extraction synchronously within this request
    intelligence_json = await _run_sarvam_extraction_sync(call.transcript)

    # Save result
    call.intelligence = intelligence_json
    call.status = (
        AnalysisStatus.COMPLETED if "error" not in intelligence_json
        else AnalysisStatus.FAILED
    )
    await db.commit()
    await db.refresh(call)

    logger.info("Intelligence extraction complete for call %s -> %s", call_id, call.status)
    response = CallAnalysisDetailResponse.model_validate(call)
    if call.audio_r2_key:
        response.playback_url = get_download_presigned_url(call.audio_r2_key, expiration=3600)
    return response


@router.get("/calls/{call_id}/export")
async def export_call(
    call_id: uuid.UUID,
    format: str = Query("pdf", description="Export format: pdf or docx"),
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a complete call analysis report as PDF or Word document."""
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    result = await db.execute(
        select(CallAnalysis).where(
            CallAnalysis.id == call_id,
            CallAnalysis.user_id == user_id,
        )
    )
    call = result.scalar_one_or_none()
    if not call:
        raise HTTPException(status_code=404, detail="Call analysis not found")

    export_data = {
        "case_id": str(call.id),
        "filename": call.filename,
        "status": call.status.value if call.status else None,
        "audio_url": call.audio_url,
        "duration_seconds": call.duration_seconds,
        "created_at": call.created_at.isoformat() if call.created_at else None,
        "intelligence": call.intelligence,
        "transcript_text": None,
    }

    # Flatten transcript into plain text
    if call.transcript:
        res_obj = call.transcript.get("result", call.transcript.get("prediction", call.transcript))
        utterances = res_obj.get("utterances") or []
        if utterances:
            lines = [f"Speaker {u.get('speaker','?')}: {u.get('text','')}" for u in utterances]
            export_data["transcript_text"] = "\n".join(lines)
        else:
            export_data["transcript_text"] = str(res_obj.get("transcription", ""))

    if format not in {"pdf", "docx"}:
        raise HTTPException(status_code=400, detail="Export format must be pdf or docx")

    import io
    import textwrap
    from fastapi.responses import StreamingResponse

    report_lines = [
        "VOICEINSIGHT CASE REPORT",
        f"Case ID: {export_data['case_id']}",
        f"Filename: {export_data['filename'] or 'N/A'}",
        f"Status: {export_data['status'] or 'N/A'}",
        f"Created: {export_data['created_at'] or 'N/A'}",
        f"Duration: {export_data['duration_seconds'] or 'N/A'} seconds",
        "",
        "INTELLIGENCE",
        json.dumps(export_data["intelligence"] or {}, indent=2, ensure_ascii=False),
        "",
        "TRANSCRIPT",
        export_data["transcript_text"] or "No transcript available.",
    ]

    if format == "docx":
        from docx import Document

        document = Document()
        document.add_heading("VoiceInsight Case Report", 0)
        for line in report_lines[1:]:
            if line in {"INTELLIGENCE", "TRANSCRIPT"}:
                document.add_heading(line.title(), level=1)
            elif line:
                document.add_paragraph(line)
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Case_{call.filename or call_id}.docx"},
        )

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    _, height = letter
    y = height - 48
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(48, y, report_lines[0])
    y -= 28
    pdf.setFont("Helvetica", 9)
    for line in "\n".join(report_lines[1:]).splitlines():
        wrapped_lines = textwrap.wrap(line, width=115) or [""]
        for wrapped_line in wrapped_lines:
            if y < 48:
                pdf.showPage()
                pdf.setFont("Helvetica", 9)
                y = height - 48
            pdf.drawString(48, y, wrapped_line)
            y -= 13
    pdf.save()
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=Case_{call.filename or call_id}.pdf"},
    )


@router.get("/analytics")
async def get_analytics(
    days: int = Query(30, ge=1, le=365, description="Lookback period in days"),
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Analytics dashboard data: threat distribution, sentiment breakdown,
    keyword frequency, calls-over-time, and emotion heatmap data.
    """
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    since = datetime.now(timezone.utc) - timedelta(days=days)

    # Fetch all completed calls in the window
    result = await db.execute(
        select(CallAnalysis).where(
            CallAnalysis.user_id == user_id,
            CallAnalysis.status == AnalysisStatus.COMPLETED,
            CallAnalysis.created_at >= since,
        )
    )
    calls = result.scalars().all()

    # Aggregate analytics
    total_calls = len(calls)
    threat_distribution = {"LOW": 0, "MEDIUM": 0, "HIGH": 0, "CRITICAL": 0}
    sentiment_distribution: dict[str, int] = {}
    keyword_frequency: dict[str, int] = {}
    calls_over_time: dict[str, int] = {}
    emotion_heatmap: dict[str, float] = {"anger": 0.0, "urgency": 0.0, "stress": 0.0, "calm": 0.0}
    languages: dict[str, int] = {}
    total_duration = 0.0

    for call in calls:
        intel = call.intelligence or {}

        # Threat
        threat = intel.get("threat_level", "LOW").upper()
        if threat in threat_distribution:
            threat_distribution[threat] += 1

        # Sentiment
        sentiment = intel.get("overall_sentiment", "Unknown")
        sentiment_distribution[sentiment] = sentiment_distribution.get(sentiment, 0) + 1

        # Keywords — support both new schema (risk_keywords_detected) and old (suspicious_keywords)
        for kw_obj in intel.get("risk_keywords_detected", []):
            kw_text = kw_obj.get("keyword", "") if isinstance(kw_obj, dict) else str(kw_obj)
            if kw_text:
                keyword_frequency[kw_text.lower()] = keyword_frequency.get(kw_text.lower(), 0) + 1
        for kw in intel.get("suspicious_keywords", []):
            kw_lower = kw.lower() if isinstance(kw, str) else str(kw).lower()
            keyword_frequency[kw_lower] = keyword_frequency.get(kw_lower, 0) + 1

        # Timeline (group by date)
        if call.created_at:
            date_key = call.created_at.strftime("%Y-%m-%d")
            calls_over_time[date_key] = calls_over_time.get(date_key, 0) + 1

        # Emotion heatmap
        emotions = intel.get("emotion_breakdown", {})
        for emo_key in emotion_heatmap:
            val = emotions.get(emo_key, 0.0)
            try:
                emotion_heatmap[emo_key] += float(val)
            except (TypeError, ValueError):
                pass

        # Language
        lang = intel.get("primary_language", "Unknown")
        languages[lang] = languages.get(lang, 0) + 1

        # Duration
        if call.duration_seconds:
            total_duration += call.duration_seconds

    # Average emotions
    if total_calls > 0:
        for k in emotion_heatmap:
            emotion_heatmap[k] = round(emotion_heatmap[k] / total_calls, 2)

    # Sort keywords by frequency (top 20)
    top_keywords = sorted(keyword_frequency.items(), key=lambda x: x[1], reverse=True)[:20]

    # Sort timeline
    sorted_timeline = sorted(calls_over_time.items())

    return {
        "period_days": days,
        "total_calls": total_calls,
        "total_duration_minutes": round(total_duration / 60, 1),
        "threat_distribution": threat_distribution,
        "sentiment_distribution": sentiment_distribution,
        "top_keywords": [{"keyword": k, "count": v} for k, v in top_keywords],
        "calls_over_time": [{"date": d, "count": c} for d, c in sorted_timeline],
        "emotion_heatmap": emotion_heatmap,
        "language_distribution": languages,
    }


@router.get("/calls/{call_id}/cross-references")
async def get_cross_references(
    call_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Multi-call cross-referencing: find other calls that share names,
    locations, or phone numbers with this call.
    """
    try:
        user_id = await clone_service.resolve_user_id(db, clerk_user_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="User not found")

    # Get target call
    result = await db.execute(
        select(CallAnalysis).where(
            CallAnalysis.id == call_id,
            CallAnalysis.user_id == user_id,
        )
    )
    target_call = result.scalar_one_or_none()
    if not target_call:
        raise HTTPException(status_code=404, detail="Call analysis not found")

    # Extract cross-reference markers from this call
    intel = target_call.intelligence or {}
    markers = intel.get("cross_reference_markers", [])
    if not markers:
        return {"cross_references": [], "marker_count": 0}

    # Build lookup sets for matching
    marker_values = set()
    for m in markers:
        normalized = m.get("normalized", m.get("value", "")).strip().lower()
        if normalized:
            marker_values.add(normalized)

    # Query all other completed calls for the same user
    result = await db.execute(
        select(CallAnalysis).where(
            CallAnalysis.user_id == user_id,
            CallAnalysis.status == AnalysisStatus.COMPLETED,
            CallAnalysis.id != call_id,
        )
    )
    other_calls = result.scalars().all()

    cross_references = []
    for other_call in other_calls:
        other_intel = other_call.intelligence or {}
        other_markers = other_intel.get("cross_reference_markers", [])

        matching_markers = []
        for om in other_markers:
            other_normalized = om.get("normalized", om.get("value", "")).strip().lower()
            if other_normalized and other_normalized in marker_values:
                matching_markers.append({
                    "type": om.get("type", "unknown"),
                    "value": om.get("value", other_normalized),
                })

        if matching_markers:
            cross_references.append({
                "call_id": str(other_call.id),
                "filename": other_call.filename,
                "created_at": other_call.created_at.isoformat() if other_call.created_at else None,
                "threat_level": other_intel.get("threat_level"),
                "matching_markers": matching_markers,
                "match_count": len(matching_markers),
            })

    # Sort by match count (most matches first)
    cross_references.sort(key=lambda x: x["match_count"], reverse=True)

    return {
        "cross_references": cross_references,
        "marker_count": len(marker_values),
        "source_markers": markers,
    }
